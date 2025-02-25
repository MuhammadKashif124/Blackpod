import os
import pandas as pd
import docx
import PyPDF2
import fitz  # PyMuPDF for better PDF image extraction
from openai import OpenAI
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import gradio as gr
from dotenv import load_dotenv
import hashlib
from PIL import Image
import io
import base64
import tempfile
import datetime
from docx import Document as DocxDocument
from pathlib import Path

# Load environment variables
load_dotenv()
api_key = os.getenv('OPENAI_API_KEY')
if not api_key:
    raise ValueError("OPENAI_API_KEY environment variable is not set.")

# Instantiate the OpenAI client
client = OpenAI(api_key=api_key)

#ollama test 


# Initialize OpenAI embeddings
embeddings = OpenAIEmbeddings(
    openai_api_key=api_key,
    model="text-embedding-3-large"
)

# Import ChromaConfig
from config import ChromaConfig

# Initialize ChromaConfig
chroma_config = ChromaConfig(use_api=os.getenv('USE_EC2', 'false').lower() == 'true')
UPLOADS_DIR = chroma_config.uploads_dir
DB_DIR = chroma_config.db_dir

# Dictionary to store ChromaDB instances
chroma_dbs = {}

def get_chroma_db(db_name):
    """Get or create a ChromaDB instance for the specified database."""
    try:
        if db_name not in chroma_dbs:
            db_path = chroma_config.get_db_path(db_name)
            # Ensure the database directory exists
            os.makedirs(db_path, exist_ok=True)
            chroma_dbs[db_name] = Chroma(
                embedding_function=embeddings,
                persist_directory=db_path
            )
        return chroma_dbs[db_name]
    except Exception as e:
        raise ValueError(f"Error loading database '{db_name}': {str(e)}")

def get_available_databases():
    """Get list of available databases."""
    return [d for d in os.listdir(DB_DIR) if os.path.isdir(os.path.join(DB_DIR, d))]

def calculate_file_hash(file_path):
    """Calculate SHA-256 hash of a file."""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()

def is_duplicate_file(file_path, db_name):
    """
    Check if file is already processed in the specific database.
    Returns True if the file exists in the specified database.
    """
    file_hash = calculate_file_hash(file_path)
    try:
        db = get_chroma_db(db_name)
        # Search for documents with matching file hash in this specific database
        results = db.similarity_search(
            query="",
            filter={"file_hash": file_hash},  # Using filter instead of search
            k=1
        )
        return len(results) > 0
    except Exception as e:
        print(f"Error checking for duplicates: {str(e)}")
        return False

def extract_images_from_pdf(pdf_path):
    """Extract images from PDF and return their descriptions."""
    image_descriptions = []
    pdf_document = fitz.open(pdf_path)
    
    for page_number in range(pdf_document.page_count):
        page = pdf_document[page_number]
        image_list = page.get_images()
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            
            # Save image temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                temp_file.write(image_bytes)
                temp_path = temp_file.name
            
            # Get image description
            description = analyze_image(temp_path)
            image_descriptions.append({
                'page': page_number + 1,
                'description': description,
                'location': f'PDF Page {page_number + 1}, Image {img_index + 1}'
            })
            
            # Clean up temporary file
            os.unlink(temp_path)
    
    pdf_document.close()
    return image_descriptions

def extract_images_from_docx(docx_path):
    """Extract images from DOCX and return their descriptions."""
    image_descriptions = []
    doc = docx.Document(docx_path)
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            # Extract image
            image_data = rel.target_part.blob
            
            # Save image temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                temp_file.write(image_data)
                temp_path = temp_file.name
            
            # Get image description
            description = analyze_image(temp_path)
            image_descriptions.append({
                'description': description,
                'location': f'Document image: {rel.target_ref}'
            })
            
            # Clean up temporary file
            os.unlink(temp_path)
    
    return image_descriptions

def analyze_image(image_path):
    """Analyze image using OpenAI's Vision model and return description."""
    # Read and encode image
    with open(image_path, "rb") as image_file:
        encoded_image = base64.b64encode(image_file.read()).decode('utf-8')
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Please provide a detailed description of this image, including any visible text, key elements, and context."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{encoded_image}",
                                "detail": "high"
                            }
                        }
                    ]
                }
            ],
            max_tokens=150
        )
        return response.choices[0].message.content
    except Exception as e:
        # Add more detailed error logging
        print(f"Error details: {str(e)}")
        return f"Error analyzing image: {str(e)}"

def process_image_file(file_path):
    """Process standalone image files."""
    try:
        # Ensure image is in a supported format
        with Image.open(file_path) as img:
            # Convert to RGB if necessary
            if img.mode != 'RGB':
                img = img.convert('RGB')
            # Save to temporary file in JPEG format
            with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_file:
                img.save(temp_file.name, 'JPEG')
                description = analyze_image(temp_file.name)
            # Clean up temporary file
            os.unlink(temp_file.name)
            return description
    except Exception as e:
        return f"Error processing image: {str(e)}"

def process_pdf(file_path):
    """Process PDF files and return text chunks with image descriptions."""
    text = ""
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    
    # Extract and process images
    image_descriptions = extract_images_from_pdf(file_path)
    
    # Add image descriptions to text
    if image_descriptions:
        text += "\n=== Images in Document ===\n"
        for img_desc in image_descriptions:
            text += f"\nLocation: {img_desc['location']}\nDescription: {img_desc['description']}\n"
    
    return text

def process_docx(file_path):
    """Process DOCX files and return text chunks with image descriptions."""
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    
    # Extract and process images
    image_descriptions = extract_images_from_docx(file_path)
    
    # Add image descriptions to text
    if image_descriptions:
        text += "\n=== Images in Document ===\n"
        for img_desc in image_descriptions:
            text += f"\nLocation: {img_desc['location']}\nDescription: {img_desc['description']}\n"
    
    return text

def process_csv(file_path):
    """Process CSV files and return text chunks."""
    df = pd.read_csv(file_path)
    return df.to_string(index=False)

def process_xlsx(file_path):
    """Process XLSX files and return text chunks."""
    df = pd.read_excel(file_path)
    return df.to_string(index=False)

def process_txt(file_path):
    """Process TXT files and return text content."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Fallback to a different encoding if UTF-8 fails
        with open(file_path, 'r', encoding='latin-1') as file:
            return file.read()

def generate_embeddings(text_chunk):
    """Generate embeddings using OpenAI API."""
    response = client.embeddings.create(
        model="text-embedding-3-large",
        input=text_chunk
    )
    return response.data[0].embedding

def generate_extra_key_data(chunk):
    """Generate extra key data using OpenAI's new API."""
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "user", "content": f"Generate a summary or key information for the following text chunk: {chunk}"}
        ]
    )
    return response.choices[0].message.content

def upsert_to_chroma_db(text_chunk, extra_key_data, document_name, db_name, file_hash):
    """Upsert embeddings into ChromaDB."""
    document = Document(
        page_content=text_chunk,
        metadata={
            "extra_key_data": extra_key_data,
            "document_name": document_name,
            "chunk_hash": hash(text_chunk),
            "file_hash": file_hash
        }
    )
    
    db = get_chroma_db(db_name)
    db.add_documents([document])

def process_file(file_path, db_name):
    """Process the file based on its type and store embeddings."""
    # Check for duplicate in the specific database
    if is_duplicate_file(file_path, db_name):
        return f"Skipped {os.path.basename(file_path)} (already exists in database '{db_name}')"

    file_hash = calculate_file_hash(file_path)
    file_extension = os.path.splitext(file_path)[1].lower()

    # Handle different file types
    if file_extension in ['.jpg', '.jpeg', '.png', '.gif']:
        # Process image file
        description = process_image_file(file_path)
        text = f"Image File: {os.path.basename(file_path)}\nDescription: {description}"
    elif file_extension == '.pdf':
        text = process_pdf(file_path)
    elif file_extension == '.docx':
        text = process_docx(file_path)
    elif file_extension == '.csv':
        text = process_csv(file_path)
    elif file_extension == '.xlsx':
        text = process_xlsx(file_path)
    elif file_extension == '.txt':  # Add support for .txt files
        text = process_txt(file_path)
    else:
        return f"Skipped {os.path.basename(file_path)} (unsupported file type)"

    # Split text into manageable chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    text_chunks = text_splitter.split_text(text)

    # Process each chunk
    for chunk in text_chunks:
        extra_key_data = generate_extra_key_data(chunk)
        upsert_to_chroma_db(chunk, extra_key_data, os.path.basename(file_path), db_name, file_hash)
    
    return f"Processed {os.path.basename(file_path)} and added to database '{db_name}'"

def upload_and_process_files(files, db_name):
    """Handle multiple file uploads and processing."""
    if not files:
        return "No files uploaded."
    
    if not db_name:
        return "Please enter a database name."
    
    # Create a directory for the specific database if it doesn't exist
    db_specific_dir = os.path.join(DB_DIR, db_name)
    os.makedirs(db_specific_dir, exist_ok=True)
    
    results = []
    for file in files:
        # Copy the file to our uploads directory with a unique name if needed
        filename = os.path.basename(file.name)
        permanent_path = os.path.join(UPLOADS_DIR, filename)
        
        # If file already exists in uploads, add a timestamp to make it unique
        if os.path.exists(permanent_path):
            base_name, ext = os.path.splitext(filename)
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{base_name}_{timestamp}{ext}"
            permanent_path = os.path.join(UPLOADS_DIR, filename)
        
        import shutil
        shutil.copy2(file.name, permanent_path)
        
        # Process the file from its new location
        result = process_file(permanent_path, db_name)
        results.append(result)
    
    return "\n".join(results)

def search_chunks(query, db_name):
    """Search for similar chunks in the specified ChromaDB."""
    db = get_chroma_db(db_name)
    # Get top 5 results
    results = db.similarity_search(
        query=query,
        k=5  # Get top 5 similar chunks
    )
    
    # Format results for display
    formatted_results = []
    for doc in results:
        formatted_results.append([
            doc.page_content,
            doc.metadata.get('extra_key_data', ''),
            doc.metadata.get('document_name', '')
        ])
    
    return formatted_results

def read_docx_file(file_path):
    """Read and extract text from DOCX file."""
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def get_file_content_as_base64(file_path):
    """Convert file content to base64 string."""
    with open(file_path, "rb") as file:
        return base64.b64encode(file.read()).decode('utf-8')

def show_preview(evt: gr.SelectData, results_df):
    """Show preview of the selected document."""
    try:
        selected_row = results_df.iloc[evt.index[0]].tolist()
        doc_name = selected_row[2]
        file_path = os.path.join(UPLOADS_DIR, doc_name)

        if not os.path.exists(file_path):
            return {
                preview_text: gr.update(visible=True, value=f"File not found: {doc_name}"),
                preview_image: gr.update(visible=False),
                preview_html: gr.update(visible=False),
                preview_table: gr.update(visible=False)
            }

        file_extension = os.path.splitext(doc_name)[1].lower()

        # Handle different file types
        if file_extension in ['.jpg', '.jpeg', '.png', '.gif']:
            return {
                preview_image: gr.update(visible=True, value=file_path),
                preview_html: gr.update(visible=False),
                preview_text: gr.update(visible=False),
                preview_table: gr.update(visible=False)
            }
        elif file_extension == '.pdf':
            # Create an HTML viewer for PDF using base64 encoding
            base64_pdf = get_file_content_as_base64(file_path)
            pdf_viewer_html = f'''
                <div style="width:100%; height:800px;">
                    <object
                        data="data:application/pdf;base64,{base64_pdf}"
                        type="application/pdf"
                        width="100%"
                        height="100%"
                    >
                        <p>Unable to display PDF file. <a href="data:application/pdf;base64,{base64_pdf}" download="{doc_name}">Download</a> instead.</p>
                    </object>
                </div>
            '''
            return {
                preview_html: gr.update(visible=True, value=pdf_viewer_html),
                preview_image: gr.update(visible=False),
                preview_text: gr.update(visible=False),
                preview_table: gr.update(visible=False)
            }
        elif file_extension in ['.docx', '.txt']:  # Add .txt to text preview
            try:
                if file_extension == '.docx':
                    text_content = read_docx_file(file_path)
                else:  # .txt file
                    with open(file_path, 'r', encoding='utf-8') as file:
                        text_content = file.read()
                return {
                    preview_text: gr.update(visible=True, value=text_content),
                    preview_image: gr.update(visible=False),
                    preview_html: gr.update(visible=False),
                    preview_table: gr.update(visible=False)
                }
            except Exception as e:
                return {
                    preview_text: gr.update(visible=True, value=f"Error reading file: {str(e)}"),
                    preview_image: gr.update(visible=False),
                    preview_html: gr.update(visible=False),
                    preview_table: gr.update(visible=False)
                }
        elif file_extension in ['.csv', '.xlsx']:
            try:
                if file_extension == '.csv':
                    df = pd.read_csv(file_path)
                else:
                    df = pd.read_excel(file_path)
                
                return {
                    preview_table: gr.update(visible=True, value=df),
                    preview_text: gr.update(visible=False),
                    preview_image: gr.update(visible=False),
                    preview_html: gr.update(visible=False)
                }
            except Exception as e:
                return {
                    preview_text: gr.update(visible=True, value=f"Error reading table file: {str(e)}"),
                    preview_table: gr.update(visible=False),
                    preview_image: gr.update(visible=False),
                    preview_html: gr.update(visible=False)
                }
        else:
            return {
                preview_text: gr.update(visible=True, value=f"Unsupported file type: {file_extension}"),
                preview_image: gr.update(visible=False),
                preview_html: gr.update(visible=False),
                preview_table: gr.update(visible=False)
            }
    except Exception as e:
        return {
            preview_text: gr.update(visible=True, value=f"Error previewing document: {str(e)}"),
            preview_image: gr.update(visible=False),
            preview_html: gr.update(visible=False),
            preview_table: gr.update(visible=False)
        }

# Gradio UI
with gr.Blocks() as demo:
    gr.Markdown("## Document Processing and Search System")

    with gr.Tab("Upload Documents"):
        file_input = gr.File(
            label="Upload Documents",
            type="filepath",
            file_count="multiple"
        )
        
        with gr.Row():
            # Dropdown for existing databases
            existing_db_dropdown = gr.Dropdown(
                label="Select or Create Database",
                choices=[""] + get_available_databases(),  # Empty choice for new database
                interactive=True,
                allow_custom_value=True,
                value=""  # Default empty value
            )
            refresh_db_button = gr.Button("🔄 Refresh")

        upload_button = gr.Button("Upload and Process")
        upload_output = gr.Textbox(label="Output", interactive=False)

        def update_db_dropdown():
            return gr.Dropdown(choices=[""] + get_available_databases())

        # Update dropdown when refresh button is clicked
        refresh_db_button.click(
            update_db_dropdown,
            outputs=[existing_db_dropdown]
        )

        # Handle file upload with selected/new database
        upload_button.click(
            upload_and_process_files,
            inputs=[file_input, existing_db_dropdown],
            outputs=upload_output
        )

    with gr.Tab("Search Chunks"):
        with gr.Row():
            # Left column for search
            with gr.Column(scale=1):
                db_dropdown = gr.Dropdown(
                    label="Select Database",
                    choices=get_available_databases(),
                    interactive=True,
                    value="" if not get_available_databases() else get_available_databases()[0]
                )
                refresh_button = gr.Button("🔄 Refresh Database List")
                search_input = gr.Textbox(label="Search Query", placeholder="Enter your search query")
                search_button = gr.Button("Search")
                search_output = gr.Dataframe(
                    label="Search Results",
                    headers=["Chunk", "Extra Key Data", "Document Name"],
                    interactive=True,  # Make it interactive to allow selection
                    wrap=True  # Enable text wrapping for better readability
                )

            # Right column for preview
            with gr.Column(scale=1):
                preview_header = gr.Markdown("## Document Preview")
                preview_image = gr.Image(
                    label="Image Preview",
                    visible=False,
                    show_label=True
                )
                preview_html = gr.HTML(  # Changed from PDF to HTML
                    label="PDF Preview",
                    visible=False
                )
                preview_text = gr.TextArea(
                    label="Text Preview",
                    visible=False,
                    show_label=True,
                    interactive=False,
                    lines=25
                )
                preview_table = gr.Dataframe(
                    label="Table Preview",
                    visible=False,
                    show_label=True,
                    wrap=True
                )

        def update_search_dropdown():
            dbs = get_available_databases()
            return gr.Dropdown(choices=dbs, value="" if not dbs else dbs[0])

        refresh_button.click(
            update_search_dropdown,
            outputs=[db_dropdown]
        )

        search_button.click(
            search_chunks,
            inputs=[search_input, db_dropdown],
            outputs=search_output
        )

        # Add event handler for search result selection
        search_output.select(
            show_preview,
            inputs=[search_output],
            outputs=[preview_image, preview_html, preview_text, preview_table]
        )

# Launch the Gradio app
if __name__ == "__main__":
    demo.launch()
